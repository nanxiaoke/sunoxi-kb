#!/usr/bin/env python3
"""
Karpathy知识库文档处理器
将raw/原始文档转换为wiki/结构化知识
集成Ollama AI处理，支持总结、分类、关键点提取等功能
"""

import os
import sys
import json
import time
import hashlib
import re
import subprocess
import argparse
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple
import requests
import frontmatter
from llm_service import LLMService
from bs4 import BeautifulSoup

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_LOG_DIR = PROJECT_ROOT / "logs"
DEFAULT_LOG_DIR.mkdir(parents=True, exist_ok=True)

# 模型自适应层
try:
    # 相对导入（当作为模块使用时）
    from .model_adapter import ModelAdapter, get_model_adapter
except ImportError:
    # 绝对导入（当作为主脚本运行时）
    try:
        from model_adapter import ModelAdapter, get_model_adapter
    except ImportError:
        # 如果都失败，创建简化版本
        logger.warning("无法导入ModelAdapter，使用简化版本")
        ModelAdapter = None
        get_model_adapter = None

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(DEFAULT_LOG_DIR / "processor.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

class OllamaClient:
    """原 Ollama API 客户端，现已重构为统一 LLM 客户端适配器"""
    
    def __init__(self, base_url: str = None, model: str = None):
        # 兼容旧接口
        self.llm = LLMService()
        self.model = model or self.llm.config.provider(self.llm.config.flow("processor_summary").providers[0]).model
            
    def _process_task(self, task_type: str, text: str, **task_kwargs) -> str:
        flow_map = {
            "summarization": "processor_summary",
            "categorization": "processor_category",
            "keypoint_extraction": "processor_keypoints",
            "entity_extraction": "entity_extraction",
        }
        flow_name = flow_map.get(task_type, "file_import_structure")
        max_chars = self.llm.config.flow(flow_name).chunk_chars
        clipped = (text or "")[:max_chars]
        prompts = {
            "summarization": f"请用中文总结下面文档，控制在{task_kwargs.get('max_length', 300)}字以内，只输出摘要。\n\n文档内容：\n{clipped}",
            "categorization": f"请将下面文档分类，只输出一个分类名称：技术、学术论文、笔记、代码、教程、新闻、其他。\n\n文档内容：\n{clipped}",
            "keypoint_extraction": f"请从下面文档提取{task_kwargs.get('num_points', 5)}个中文关键点，使用编号列表，只输出关键点。\n\n文档内容：\n{clipped}",
            "entity_extraction": f"请从下面文档提取核心实体、术语、产品名、模型名和项目名，用逗号分隔，只输出实体列表。\n\n文档内容：\n{clipped}",
        }
        messages = [
            {"role": "system", "content": "你是技术知识库文档处理助手，严格按要求输出，不添加解释。"},
            {"role": "user", "content": prompts.get(task_type, clipped)},
        ]
        try:
            result = self.llm.chat(flow_name, messages, options={"think": False, "temperature": 0.2, "num_predict": 800})
            if result.status != "ok":
                logger.error(f"LLM 任务执行失败 ({task_type}): {result.error}")
                return None
            self.model = result.model
            logger.info(
                "Processor LLM provider: %s / %s%s",
                result.provider,
                result.model,
                f" (fallback from {result.fallback_from})" if result.fallback_from else "",
            )
            return result.content
        except Exception as e:
            logger.error(f"LLM 任务执行失败 ({task_type}): {e}")
            return None

    def generate(self, prompt: str, options=None, max_retries=3) -> str:
        result = self.llm.chat("file_import_structure", [{"role": "user", "content": prompt}], options=options or {})
        if result.status != "ok":
            logger.error(f"LLM generate失败: {result.error}")
            return None
        self.model = result.model
        return result.content
        
    def chat(self, messages: list, options=None, max_retries=3) -> str:
        result = self.llm.chat("qa", messages, options=options or {})
        if result.status != "ok":
            logger.error(f"LLM chat失败: {result.error}")
            return None
        self.model = result.model
        return result.content

    def summarize(self, text: str, max_length: int = 300) -> str:
        truncated_text = text[:5000] if len(text) > 5000 else text
        return self._process_task("summarization", truncated_text, max_length=max_length)
        
    def categorize(self, text: str) -> str:
        truncated_text = text[:3000] if len(text) > 3000 else text
        return self._process_task("categorization", truncated_text)
        
    def extract_keypoints(self, text: str, num_points: int = 5) -> str:
        truncated_text = text[:5000] if len(text) > 5000 else text
        return self._process_task("keypoint_extraction", truncated_text, num_points=num_points)
        
    def extract_entities(self, text: str) -> str:
        truncated_text = text[:3000] if len(text) > 3000 else text
        return self._process_task("entity_extraction", truncated_text)

class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self, ollama_client: Optional[OllamaClient] = None):
        self.ollama = ollama_client or OllamaClient()
        self.base_dir = PROJECT_ROOT
        
        # 输出目录
        self.wiki_dir = self.base_dir / "wiki"
        self.outputs_dir = self.base_dir / "outputs"
        
        # 创建目录
        self.wiki_dir.mkdir(parents=True, exist_ok=True)
        self.outputs_dir.mkdir(parents=True, exist_ok=True)
    
    def read_document(self, filepath: Path) -> Optional[Dict[str, Any]]:
        """读取文档内容"""
        try:
            if not filepath.exists():
                logger.error(f"文件不存在: {filepath}")
                return None
            
            # 根据文件类型选择读取方式
            if filepath.suffix.lower() in ['.md', '.markdown']:
                return self._read_markdown(filepath)
            elif filepath.suffix.lower() in ['.txt', '.text']:
                return self._read_text(filepath)
            elif filepath.suffix.lower() in ['.py', '.js', '.java', '.cpp', '.c', '.go', '.rs']:
                return self._read_code(filepath)
            elif filepath.suffix.lower() == '.pdf':
                return self._read_pdf(filepath)
            elif filepath.suffix.lower() == '.docx':
                return self._read_docx(filepath)
            else:
                # 默认按文本读取
                return self._read_text(filepath)
                
        except Exception as e:
            logger.error(f"读取文档失败 {filepath}: {e}")
            return None
    
    def _read_markdown(self, filepath: Path) -> Dict[str, Any]:
        """读取Markdown文档"""
        content = filepath.read_text(encoding='utf-8')
        
        # 解析frontmatter（如果有）
        try:
            post = frontmatter.loads(content)
            metadata = post.metadata
            content_text = post.content
        except:
            metadata = {}
            content_text = content
        
        # 提取标题
        title = metadata.get('title', '')
        if not title:
            # 从第一行标题提取
            lines = content_text.split('\n')
            for line in lines[:5]:
                if line.startswith('# '):
                    title = line[2:].strip()
                    break
        
        if not title:
            title = filepath.stem
        
        return {
            "filepath": str(filepath),
            "title": title,
            "content": content_text,
            "metadata": metadata,
            "format": "markdown",
            "hash": hashlib.md5(content_text.encode()).hexdigest()[:8]
        }
    
    def _read_text(self, filepath: Path) -> Dict[str, Any]:
        """读取纯文本文档"""
        content = filepath.read_text(encoding='utf-8')
        
        # 提取标题（第一行或文件名）
        lines = content.split('\n')
        title = filepath.stem
        
        # 如果第一行看起来像标题，使用它
        if lines and len(lines[0].strip()) < 100 and not lines[0].strip().startswith('#'):
            title = lines[0].strip()
        
        return {
            "filepath": str(filepath),
            "title": title,
            "content": content,
            "metadata": {},
            "format": "text",
            "hash": hashlib.md5(content.encode()).hexdigest()[:8]
        }
    
    def _read_code(self, filepath: Path) -> Dict[str, Any]:
        """读取代码文件"""
        content = filepath.read_text(encoding='utf-8')
        
        # 从注释中提取描述
        description = ""
        lines = content.split('\n')
        for line in lines[:10]:
            line = line.strip()
            if line.startswith('#') or line.startswith('//') or line.startswith('/*'):
                # 移除注释符号
                clean_line = line.lstrip('#/ *')
                if clean_line and len(description) < 200:
                    description += clean_line + " "
        
        title = filepath.stem
        if description:
            title = f"{title} - {description[:50]}..."
        
        return {
            "filepath": str(filepath),
            "title": title,
            "content": content,
            "metadata": {"language": filepath.suffix[1:]},
            "format": "code",
            "hash": hashlib.md5(content.encode()).hexdigest()[:8]
        }
    
    def _read_pdf(self, filepath: Path) -> Dict[str, Any]:
        """读取PDF文档"""
        try:
            import pdfplumber
            content_parts = []
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content_parts.append(text)
            content = '\n\n'.join(content_parts)
        except ImportError:
            logger.warning("pdfplumber未安装，尝试pdfminer降级...")
            try:
                from pdfminer.high_level import extract_text
                content = extract_text(str(filepath))
            except ImportError:
                logger.error("没有可用的PDF解析库，请安装pdfplumber")
                content = f"[PDF文件需安装pdfplumber解析: {filepath.name}]"
        except Exception as e:
            logger.error(f"PDF解析失败 {filepath}: {e}")
            content = f"[PDF解析失败: {filepath.name}]"

        if not content.strip():
            content = f"[PDF未提取到文本: {filepath.name}]"

        return {
            "filepath": str(filepath),
            "title": filepath.stem,
            "content": content,
            "metadata": {"pages": len(content_parts) if 'content_parts' in dir() else 0},
            "format": "pdf",
            "hash": hashlib.md5(content.encode()).hexdigest()[:8]
        }

    def _read_docx(self, filepath: Path) -> Dict[str, Any]:
        """读取Word文档"""
        try:
            from docx import Document
            doc = Document(str(filepath))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = '\n\n'.join(paragraphs)
            # 提取标题（从文档样式或文件名）
            title = filepath.stem
            for p in doc.paragraphs:
                if p.style and 'heading' in (p.style.name or '').lower():
                    title = p.text.strip()
                    break
        except ImportError:
            logger.error("python-docx未安装，请安装python-docx")
            content = f"[Word文件需安装python-docx解析: {filepath.name}]"
            title = filepath.stem
        except Exception as e:
            logger.error(f"DOCX解析失败 {filepath}: {e}")
            content = f"[DOCX解析失败: {filepath.name}]"
            title = filepath.stem

        return {
            "filepath": str(filepath),
            "title": title,
            "content": content,
            "metadata": {},
            "format": "docx",
            "hash": hashlib.md5(content.encode()).hexdigest()[:8]
        }

    def process_document(self, doc_info: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """处理单个文档，具有强健的错误处理和降级策略"""
        try:
            logger.info(f"处理文档: {doc_info['title']}")
            
            # 初始化结果
            summary = "（AI生成摘要失败）"
            keypoints = "（关键点提取失败）"
            metadata = doc_info.get("metadata") or {}
            category_override = (metadata.get("preferred_category") or metadata.get("category") or "").strip()
            category = category_override or "其他"
            entities = ""
            model_used = "降级处理（无AI）"
            
            # 1. 尝试生成摘要
            try:
                logger.debug(f"开始生成摘要: {doc_info['title']}")
                summary_result = self.ollama.summarize(doc_info["content"])
                if summary_result and summary_result.strip():
                    summary = summary_result.strip()
                    logger.debug(f"摘要生成完成: {len(summary)}字符")
                    model_used = self.ollama.model
                else:
                    logger.warning(f"生成摘要失败或返回空值: {doc_info['title']}")
                    # 降级处理：使用前100个字符作为摘要
                    content = doc_info["content"]
                    if len(content) > 100:
                        summary = content[:100] + "..."
                    else:
                        summary = content
            except Exception as e:
                logger.warning(f"生成摘要异常: {e}")
                # 降级处理
                content = doc_info["content"]
                if len(content) > 100:
                    summary = content[:100] + "..."
                else:
                    summary = content
            
            # 2. 尝试提取关键点
            try:
                logger.debug(f"开始提取关键点: {doc_info['title']}")
                keypoints_result = self.ollama.extract_keypoints(doc_info["content"])
                if keypoints_result and keypoints_result.strip():
                    keypoints = keypoints_result.strip()
                    logger.debug(f"关键点提取完成: {len(keypoints)}字符")
                else:
                    logger.warning(f"提取关键点失败或返回空值: {doc_info['title']}")
                    # 降级处理：使用前几个段落作为关键点
                    lines = [line.strip() for line in doc_info["content"].split('\n') if line.strip()]
                    if lines:
                        keypoints = "\n".join([f"{i+1}. {line[:80]}..." for i, line in enumerate(lines[:3])])
            except Exception as e:
                logger.warning(f"提取关键点异常: {e}")
                # 降级处理
                lines = [line.strip() for line in doc_info["content"].split('\n') if line.strip()]
                if lines:
                    keypoints = "\n".join([f"{i+1}. {line[:80]}..." for i, line in enumerate(lines[:3])])
            
            # 3. 尝试分类；若审核阶段已指定分类，则直接使用。
            if category_override:
                logger.info(f"使用审核指定分类: {doc_info['title']} -> {category_override}")
            else:
                try:
                    logger.debug(f"开始分类: {doc_info['title']}")
                    category_result = self.ollama.categorize(doc_info["content"])
                    if category_result and category_result.strip():
                        category = category_result.strip()
                        logger.debug(f"分类完成: {category}")
                    else:
                        logger.warning(f"分类失败或返回空值: {doc_info['title']}")
                        # 降级处理：基于文件扩展名或内容猜测
                        if doc_info.get("format") == "code":
                            category = "代码"
                        elif doc_info.get("format") == "markdown":
                            category = "笔记"
                        else:
                            category = "其他"
                except Exception as e:
                    logger.warning(f"分类异常: {e}")
                    # 降级处理
                    if doc_info.get("format") == "code":
                        category = "代码"
                    elif doc_info.get("format") == "markdown":
                        category = "笔记"
                    else:
                        category = "其他"
            
            # 4. 尝试提取实体
            try:
                logger.debug(f"开始提取实体: {doc_info['title']}")
                entities_result = self.ollama.extract_entities(doc_info["content"])
                if entities_result and entities_result.strip():
                    entities = entities_result.strip()
                    logger.debug(f"实体提取完成: {len(entities.split())}个实体")
                else:
                    logger.debug(f"未提取到实体: {doc_info['title']}")
            except Exception as e:
                logger.warning(f"提取实体异常: {e}")
                # 实体提取失败时留空
                entities = ""
            
            # 构建结果
            result = {
                **doc_info,
                "summary": summary,
                "keypoints": keypoints,
                "category": category.strip(),
                "entities": entities,
                "tags": metadata.get("tags") or [],
                "processed_at": datetime.now().isoformat(),
                "model_used": model_used
            }
            
            logger.info(f"文档处理完成: {doc_info['title']} -> {category} (模型: {model_used})")
            return result
            
        except Exception as e:
            logger.error(f"处理文档失败 {doc_info.get('title', '未知')}: {e}", exc_info=True)
            return None
    
    def save_to_wiki(self, result: Dict[str, Any]) -> Optional[Path]:
        """保存处理结果到wiki目录"""
        try:
            # 确定分类目录
            category = result["category"]
            category_dir = self.wiki_dir / "concepts"  # 默认
            
            # 映射分类到目录
            category_map = {
                "技术": "technologies",
                "学术论文": "concepts",
                "笔记": "notes",
                "代码": "technologies",
                "文章": "concepts",
                "新闻": "concepts",
                "教程": "technologies",
                "其他": "concepts"
            }
            
            wiki_category = category_map.get(category, "concepts")
            category_dir = self.wiki_dir / wiki_category
            category_dir.mkdir(parents=True, exist_ok=True)
            
            # 生成文件名
            title_slug = "".join(c for c in result["title"] if c.isalnum() or c in " -_")[:50]
            filename = f"{title_slug}_{result['hash']}.md"
            filepath = category_dir / filename
            
            result_tags = result.get("tags") or []
            if isinstance(result_tags, str):
                result_tags = [t.strip() for t in re.split(r"[,，;；]", result_tags) if t.strip()]
            tag_lines = "\n".join([f"  - {t}" for t in ([result['category']] + [t for t in result_tags if t != result['category']])])
            # 构建wiki内容
            wiki_content = f"""---
title: "{result['title']}"
category: {result['category']}
date: {result['processed_at']}
model: {result['model_used']}
source: {result['filepath']}
tags:
{tag_lines}
---

# {result['title']}

> **来源**: {result['filepath']}
> **处理时间**: {result['processed_at']}
> **分类**: {result['category']}
> **AI模型**: {result['model_used']}
> **原始格式**: {result['format']}

## 📝 摘要

{result['summary']}

## 🔑 关键点

{result['keypoints']}

## 🏷️ 实体与概念

{result['entities'] if result['entities'] else "（未提取到实体）"}

## 📄 原始内容

> 注意：这是AI处理的摘要版本，查看完整内容请访问原始文件。

{result['content'][:1000]}...
"""
            
            # 保存文件
            filepath.write_text(wiki_content, encoding='utf-8')
            logger.info(f"Wiki保存成功: {filepath}")
            
            # 同时保存JSON格式的完整结果（用于后续处理）
            json_path = self.outputs_dir / f"{result['hash']}.json"
            json_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding='utf-8')
            
            return filepath
            
        except Exception as e:
            logger.error(f"保存Wiki失败: {e}")
            return None
    
    def _save_progress(self, processed_files: set) -> None:
        """保存处理进度"""
        try:
            progress_file = self.base_dir / "progress.json"
            progress_data = {
                "processed_files": list(processed_files),
                "saved_at": datetime.now().isoformat()
            }
            progress_file.write_text(json.dumps(progress_data, ensure_ascii=False, indent=2), encoding='utf-8')
            logger.debug(f"进度已保存: {len(processed_files)} 个文件")
        except Exception as e:
            logger.warning(f"保存进度失败: {e}")
    
    def _process_with_timeout(self, doc_info: Dict[str, Any], timeout_seconds: int = 300) -> Optional[Dict[str, Any]]:
        """带超时控制的文档处理"""
        import threading
        import queue
        
        result_queue = queue.Queue()
        
        def process_worker():
            """工作线程"""
            try:
                result = self.process_document(doc_info)
                result_queue.put((True, result))
            except Exception as e:
                result_queue.put((False, str(e)))
        
        # 启动工作线程
        worker_thread = threading.Thread(target=process_worker)
        worker_thread.daemon = True  # 设置为守护线程
        worker_thread.start()
        
        # 等待结果或超时
        try:
            success, result = result_queue.get(timeout=timeout_seconds)
            if success:
                return result
            else:
                logger.error(f"工作线程异常: {result}")
                return None
        except queue.Empty:
            logger.warning(f"文档处理超时 ({timeout_seconds}秒): {doc_info['title']}")
            # 尝试强制中断（尽管守护线程会在主线程结束时退出）
            return None
        except Exception as e:
            logger.error(f"处理超时控制异常: {e}")
            return None
    
    def batch_process(self, input_dir: Path, pattern: str = "**/*") -> List[Dict[str, Any]]:
        """批量处理目录中的文档，具有断点续传和超时控制"""
        results = []
        
        # 查找文件
        files = list(input_dir.glob(pattern))
        files = [f for f in files if f.is_file() and not f.name.startswith('.')]
        
        if not files:
            logger.warning(f"未找到文件: {input_dir}/{pattern}")
            return results
        
        logger.info(f"找到 {len(files)} 个文件，开始批量处理...")
        
        # 断点续传：检查已处理的文件
        progress_file = self.base_dir / "progress.json"
        processed_files = set()
        
        if progress_file.exists():
            try:
                progress_data = json.loads(progress_file.read_text(encoding='utf-8'))
                processed_files = set(progress_data.get("processed_files", []))
                logger.info(f"断点续传: 跳过 {len(processed_files)} 个已处理的文件")
            except Exception as e:
                logger.warning(f"读取进度文件失败: {e}")
        
        # 过滤已处理的文件
        files_to_process = [f for f in files if str(f) not in processed_files]
        if not files_to_process:
            logger.info(f"所有文件已处理完成")
            return results
        
        logger.info(f"实际需要处理: {len(files_to_process)} 个文件")
        
        for i, filepath in enumerate(files_to_process):
            logger.info(f"处理进度: {i+1}/{len(files_to_process)} - {filepath.name}")
            
            try:
                # 读取文档
                doc_info = self.read_document(filepath)
                if not doc_info:
                    logger.warning(f"跳过无法读取的文件: {filepath}")
                    processed_files.add(str(filepath))
                    self._save_progress(processed_files)
                    continue
                
                # 处理文档（带超时保护）
                result = self._process_with_timeout(doc_info, timeout_seconds=300)  # 5分钟超时
                if not result:
                    logger.warning(f"处理失败: {filepath}")
                    processed_files.add(str(filepath))
                    self._save_progress(processed_files)
                    continue
                
                # 保存到wiki
                wiki_path = self.save_to_wiki(result)
                if wiki_path:
                    results.append({
                        "input": str(filepath),
                        "output": str(wiki_path),
                        "title": result["title"],
                        "category": result["category"],
                        "success": True
                    })
                    logger.info(f"✅ 处理成功: {filepath.name}")
                else:
                    results.append({
                        "input": str(filepath),
                        "output": None,
                        "title": result["title"],
                        "category": result["category"],
                        "success": False
                    })
                    logger.warning(f"⚠️ 保存失败: {filepath.name}")
                
                # 标记为已处理
                processed_files.add(str(filepath))
                self._save_progress(processed_files)
                
            except Exception as e:
                logger.error(f"批量处理异常 {filepath.name}: {e}", exc_info=True)
                results.append({
                    "input": str(filepath),
                    "output": None,
                    "title": filepath.name,
                    "category": "错误",
                    "success": False,
                    "error": str(e)
                })
                logger.warning(f"❌ 处理异常，继续下一个文件: {filepath.name}")
                
                # 标记为已处理（即使是失败）
                processed_files.add(str(filepath))
                self._save_progress(processed_files)
            
            # 延迟，避免过快请求
            if i < len(files_to_process) - 1:
                logger.debug(f"处理完成，延迟1秒后继续...")
                time.sleep(1)
        
        # 处理完成，清理进度文件
        try:
            if progress_file.exists():
                progress_file.unlink()
                logger.debug("进度文件已清理")
        except Exception as e:
            logger.warning(f"清理进度文件失败: {e}")
        
        # 统计结果
        success_count = sum(1 for r in results if r.get("success", False))
        failure_count = len(results) - success_count
        logger.info(f"批量处理完成: {success_count}成功, {failure_count}失败")
        
        return results

    # ====== AI方法代理 ======
    
    def summarize(self, text: str, max_length: int = 300) -> Optional[str]:
        """总结文本（代理方法）"""
        try:
            return self.ollama.summarize(text, max_length)
        except Exception as e:
            logger.error(f"总结失败: {e}")
            return None
    
    def categorize(self, text: str) -> Optional[str]:
        """对文档进行分类（代理方法）"""
        try:
            return self.ollama.categorize(text)
        except Exception as e:
            logger.error(f"分类失败: {e}")
            return None
    
    def extract_keypoints(self, text: str, num_points: int = 5) -> Optional[str]:
        """提取关键点（代理方法）"""
        try:
            return self.ollama.extract_keypoints(text, num_points)
        except Exception as e:
            logger.error(f"提取关键点失败: {e}")
            return None
    
    def extract_entities(self, text: str) -> Optional[str]:
        """提取实体（代理方法）"""
        try:
            return self.ollama.extract_entities(text)
        except Exception as e:
            logger.error(f"提取实体失败: {e}")
            return None


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description="Karpathy知识库文档处理器")
    parser.add_argument("--input", "-i", help="输入目录或文件")
    parser.add_argument("--output", "-o", help="输出目录（默认: wiki/）")
    parser.add_argument("--pattern", "-p", default="**/*", help="文件匹配模式（默认: **/*）")
    parser.add_argument("--model", "-m", default="gemma4:e4b", help="Ollama模型（默认: gemma4:e4b）")
    parser.add_argument("--test", action="store_true", help="测试模式")
    
    args = parser.parse_args()
    
    # 测试模式
    if args.test:
        logger.info("运行测试模式...")
        
        # 初始化客户端
        logger.info("初始化Ollama客户端...")
        try:
            ollama = OllamaClient(model=args.model)
            logger.info(f"Ollama客户端初始化完成，使用模型: {ollama.model}")
        except Exception as e:
            logger.error(f"初始化Ollama客户端失败: {e}")
            return 1
        
        # 测试连接
        test_prompt = "What is artificial intelligence in one sentence?"
        logger.info(f"发送测试请求: {test_prompt}")
        
        # 使用简单选项测试
        response = ollama.generate(test_prompt, options={"num_predict": 50, "temperature": 0.7})
        
        if response:
            logger.info(f"✅ 测试成功！Ollama响应: {response[:100]}...")
            return 0
        else:
            logger.error("❌ 测试失败：Ollama无响应")
            return 1
    
    # 正常处理模式
    if not args.input:
        parser.print_help()
        return 1
    
    input_path = Path(args.input)
    if not input_path.exists():
        logger.error(f"输入路径不存在: {input_path}")
        return 1
    
    # 初始化处理器
    processor = DocumentProcessor(OllamaClient(model=args.model))
    
    # 处理单个文件或目录
    if input_path.is_file():
        logger.info(f"处理单个文件: {input_path}")
        
        # 读取文档
        doc_info = processor.read_document(input_path)
        if not doc_info:
            logger.error(f"无法读取文件: {input_path}")
            return 1
        
        # 处理文档
        result = processor.process_document(doc_info)
        if not result:
            logger.error(f"处理失败: {input_path}")
            return 1
        
        # 保存到wiki
        wiki_path = processor.save_to_wiki(result)
        if wiki_path:
            logger.info(f"✅ 处理完成！保存到: {wiki_path}")
            return 0
        else:
            logger.error(f"❌ 保存失败")
            return 1
    
    else:  # 目录
        logger.info(f"批量处理目录: {input_path}")
        
        results = processor.batch_process(input_path, args.pattern)
        
        # 统计结果
        success_count = sum(1 for r in results if r["success"])
        total_count = len(results)
        
        logger.info(f"处理完成: {success_count}/{total_count} 成功")
        
        # 输出详细结果
        for result in results:
            status = "✅ 成功" if result["success"] else "❌ 失败"
            logger.info(f"  {status}: {result['title']} -> {result.get('category', '未知')}")
        
        return 0 if success_count > 0 else 1

if __name__ == "__main__":
    sys.exit(main())
